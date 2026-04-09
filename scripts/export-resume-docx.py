#!/usr/bin/env python3
"""
Export a structured resume markdown file to a two-column .docx
matching the JFO resume format (dark sidebar + main content).

Usage:
    python3 scripts/export-resume-docx.py <input.md> [output.docx]

If output path is omitted, writes to ~/Desktop/ with the input filename.
"""

import sys
import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Brand colors ──────────────────────────────────────────────
ORANGE = RGBColor(0xE8, 0x53, 0x0E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_BG = "333333"
LIGHT_GRAY = RGBColor(0x55, 0x55, 0x55)
FONT_NAME = "Calibri"

# ── Layout constants ─────────────────────────────────────────
LEFT_COL_WIDTH = Inches(2.25)
RIGHT_COL_WIDTH = Inches(5.25)
PAGE_TOP_MARGIN = Inches(0.3)
PAGE_BOTTOM_MARGIN = Inches(0.25)
PAGE_LEFT_MARGIN = Inches(0.3)
PAGE_RIGHT_MARGIN = Inches(0.3)


# ── Helpers ───────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_para(cell, text="", font_size=9, color=BLACK, bold=False, italic=False,
             space_before=0, space_after=1, alignment=None, first=False):
    """Add a paragraph to a cell with styling."""
    if first:
        p = cell.paragraphs[0]
        p.text = ""
    else:
        p = cell.add_paragraph()

    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic

    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = Pt(font_size + 1)
    if alignment:
        fmt.alignment = alignment

    return p


def add_separator(cell, color=ORANGE):
    """Add a thin horizontal rule / separator line."""
    p = cell.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(1)
    fmt.space_after = Pt(3)
    # Use bottom border on paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{color_to_hex(color)}"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def color_to_hex(rgb_color):
    return f"{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}" if isinstance(rgb_color, tuple) else str(rgb_color)


# ── Parsing ──────────────────────────────────────────────────

def parse_resume_md(filepath):
    """
    Parse the structured markdown into sidebar and main content.
    Split point: line containing the candidate's name (all-caps, multi-word line
    after the sidebar skill blocks).
    """
    with open(filepath, "r") as f:
        lines = f.readlines()

    lines = [l.rstrip("\n") for l in lines]

    # Find the split point: first all-caps line with 2+ words that isn't a
    # sidebar section header (sidebar headers are followed by ■ lines)
    split_idx = None
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        # Check if this is an all-caps multi-word line
        words = stripped.split()
        if (len(words) >= 2
                and stripped == stripped.upper()
                and not stripped.startswith("■")
                and any(c.isalpha() for c in stripped)):
            # Look ahead — if next non-empty line starts with ■, this is a sidebar header
            is_sidebar_header = False
            for j in range(i + 1, min(i + 3, len(lines))):
                nxt = lines[j].strip()
                if not nxt:
                    continue
                if nxt.startswith("■"):
                    is_sidebar_header = True
                break
            if not is_sidebar_header:
                # Check it's a name-like line (contains letters, not a section like EXPERIENCE)
                # Names typically come after sidebar content
                if i > 5:  # must be past the contact block at minimum
                    split_idx = i
                    break

    if split_idx is None:
        # Fallback: look for a line that looks like a name
        for i, line in enumerate(lines):
            if "O'CONNOR" in line.upper() or "O'CONNOR" in line:
                split_idx = i
                break

    if split_idx is None:
        print("Warning: Could not find name line to split sidebar/main. Using line 45.")
        split_idx = 45

    sidebar_lines = lines[:split_idx]
    main_lines = lines[split_idx:]

    return sidebar_lines, main_lines


def parse_sidebar_blocks(lines):
    """
    Parse sidebar lines into blocks:
    [{'header': 'CONTACT', 'items': ['hi@jim.software', ...]}, ...]
    Non-header, non-■ lines in the CONTACT block are plain contact items.
    ■ lines are skill items.
    """
    blocks = []
    current = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("■"):
            if current:
                current["items"].append(stripped)
            continue

        # Check if it's a section header (all caps, no ■)
        if stripped == stripped.upper() and len(stripped) > 1 and any(c.isalpha() for c in stripped):
            if current:
                blocks.append(current)
            current = {"header": stripped, "items": []}
        else:
            # Plain text line (contact info, etc.)
            if current:
                current["items"].append(stripped)

    if current:
        blocks.append(current)

    return blocks


def parse_main_content(lines):
    """
    Parse main content lines into structured elements:
    - name, tagline, separator, summary
    - sections (EXPERIENCE, EDUCATION, 0→1 PLATFORMS)
    - within sections: company blocks, sub-sections, bullets
    """
    elements = []
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Name (first non-empty line)
        if not elements:
            elements.append({"type": "name", "text": line})
            i += 1
            continue

        # Tagline (second element, after name)
        if len(elements) == 1:
            elements.append({"type": "tagline", "text": line})
            i += 1
            continue

        # Separator
        if line.startswith("———") or line.startswith("---"):
            elements.append({"type": "separator"})
            i += 1
            continue

        # Section headers (all caps)
        if (line == line.upper()
                and len(line) > 2
                and any(c.isalpha() for c in line)
                and not line.startswith("—")):
            elements.append({"type": "section_header", "text": line})
            i += 1
            continue

        # Company line (contains year range pattern like "2021 – Present")
        if ("–" in line or "–" in line) and not line.startswith("—"):
            # Could be "SAP Concur    2021 – Present · Allen, TX"
            # or "University of Cambridge — PhD Candidate..."
            if any(y in line for y in ["2017", "2018", "2019", "2020", "2021", "2022", "2023", "2024", "2025", "2026", "Present"]):
                # Check if it's an education line (contains — as separator within text)
                if line.startswith("University") or line.startswith("Newline") or line.startswith("—"):
                    elements.append({"type": "education_line", "text": line})
                else:
                    elements.append({"type": "company", "text": line})
                i += 1
                continue

        # Role line (italic — follows company line)
        if elements and elements[-1]["type"] == "company" and not line.startswith("—"):
            elements.append({"type": "role", "text": line})
            i += 1
            continue

        # Sub-section header (bold, has date range in parens)
        if ("(" in line and ")" in line and
                any(y in line for y in ["2021", "2022", "2023", "2024", "2025", "Present"])):
            elements.append({"type": "subsection", "text": line})
            i += 1
            continue

        # Bullet points (em-dash prefix)
        if line.startswith("—"):
            elements.append({"type": "bullet", "text": line})
            i += 1
            continue

        # Project header (contains — as separator, like "Frame OS — frame.jim.software")
        if " — " in line and not line.startswith("—"):
            elements.append({"type": "project_header", "text": line})
            i += 1
            continue

        # Education / training lines
        if any(kw in line for kw in ["University", "Newline", "CCNA", "iOS Developer", "PADI", "Boot Camp", "Bootcamp"]):
            elements.append({"type": "education_line", "text": line})
            i += 1
            continue

        # Default: body text
        elements.append({"type": "body", "text": line})
        i += 1

    return elements


# ── Document generation ──────────────────────────────────────

def build_docx(sidebar_blocks, main_elements, output_path):
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = PAGE_TOP_MARGIN
    section.bottom_margin = PAGE_BOTTOM_MARGIN
    section.left_margin = PAGE_LEFT_MARGIN
    section.right_margin = PAGE_RIGHT_MARGIN

    # ── Create two-column table ──
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    remove_table_borders(table)

    # Set column widths
    table.columns[0].width = LEFT_COL_WIDTH
    table.columns[1].width = RIGHT_COL_WIDTH

    # Also set cell widths explicitly
    for row in table.rows:
        row.cells[0].width = LEFT_COL_WIDTH
        row.cells[1].width = RIGHT_COL_WIDTH

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    # ── Dark sidebar background ──
    set_cell_shading(left_cell, DARK_BG)

    # Add cell padding
    for cell, pad_left, pad_right in [(left_cell, 200, 120), (right_cell, 200, 80)]:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'  <w:top w:w="100" w:type="dxa"/>'
            f'  <w:left w:w="{pad_left}" w:type="dxa"/>'
            f'  <w:bottom w:w="100" w:type="dxa"/>'
            f'  <w:right w:w="{pad_right}" w:type="dxa"/>'
            '</w:tcMar>'
        )
        tcPr.append(tcMar)

    # ── Populate LEFT (sidebar) ──
    first_para = True
    for block in sidebar_blocks:
        header = block["header"]
        items = block["items"]

        # Section header
        add_para(left_cell, header,
                 font_size=8, color=ORANGE, bold=True,
                 space_before=3 if not first_para else 0,
                 space_after=0,
                 first=first_para)
        first_para = False

        # Underline separator for header
        add_separator_line(left_cell)

        # Items
        for item in items:
            text = item.lstrip("■ ").strip()
            is_skill = item.strip().startswith("■")

            if is_skill:
                p = left_cell.add_paragraph()
                # Add the bullet character
                run_bullet = p.add_run("■ ")
                run_bullet.font.name = FONT_NAME
                run_bullet.font.size = Pt(4)
                run_bullet.font.color.rgb = ORANGE

                run_text = p.add_run(text)
                run_text.font.name = FONT_NAME
                run_text.font.size = Pt(8)
                run_text.font.color.rgb = WHITE

                fmt = p.paragraph_format
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0.5)
                fmt.line_spacing = Pt(9)
            else:
                # Contact info or plain text
                is_project_desc = header == "PROJECTS & RESEARCH" and not text.startswith(("hi@", "832", "McKinney", "jim.", "github"))
                if is_project_desc:
                    add_para(left_cell, text,
                             font_size=6.5, color=LIGHT_GRAY,
                             space_before=0, space_after=3, italic=True)
                else:
                    add_para(left_cell, text,
                             font_size=8, color=WHITE,
                             space_before=0, space_after=1)

    # ── Populate RIGHT (main content) ──
    first_right = True
    for el in main_elements:
        t = el["type"]

        if t == "name":
            add_para(right_cell, el["text"],
                     font_size=20, color=BLACK, bold=True,
                     space_before=0, space_after=1,
                     first=first_right)
            first_right = False

        elif t == "tagline":
            add_para(right_cell, el["text"],
                     font_size=9, color=ORANGE, bold=True, italic=True,
                     space_before=0, space_after=1)

        elif t == "separator":
            add_separator_line(right_cell, ORANGE)

        elif t == "body":
            add_para(right_cell, el["text"],
                     font_size=8, color=BLACK,
                     space_before=0, space_after=2)

        elif t == "section_header":
            add_para(right_cell, el["text"],
                     font_size=10, color=ORANGE, bold=True,
                     space_before=3, space_after=1)

        elif t == "company":
            text = el["text"]
            # Split company name from date/location if possible
            # Pattern: "SAP Concur                    2021 – Present · Allen, TX"
            parts = text.split("  ")  # multiple spaces as separator
            parts = [p.strip() for p in parts if p.strip()]

            p = right_cell.add_paragraph()
            if len(parts) >= 2:
                company_name = parts[0]
                date_loc = parts[-1]

                run_co = p.add_run(company_name)
                run_co.font.name = FONT_NAME
                run_co.font.size = Pt(10)
                run_co.font.color.rgb = BLACK
                run_co.font.bold = True

                # Add tab + date/location right-aligned
                run_sep = p.add_run("    ")
                run_sep.font.size = Pt(8)

                run_date = p.add_run(date_loc)
                run_date.font.name = FONT_NAME
                run_date.font.size = Pt(8)
                run_date.font.color.rgb = LIGHT_GRAY
            else:
                run = p.add_run(text)
                run.font.name = FONT_NAME
                run.font.size = Pt(10)
                run.font.color.rgb = BLACK
                run.font.bold = True

            fmt = p.paragraph_format
            fmt.space_before = Pt(2)
            fmt.space_after = Pt(0)

        elif t == "role":
            add_para(right_cell, el["text"],
                     font_size=8, color=LIGHT_GRAY, italic=True,
                     space_before=0, space_after=1)

        elif t == "subsection":
            add_para(right_cell, el["text"],
                     font_size=9, color=BLACK, bold=True,
                     space_before=2, space_after=0)

        elif t == "project_header":
            add_para(right_cell, el["text"],
                     font_size=9, color=BLACK, bold=True,
                     space_before=2, space_after=0)

        elif t == "bullet":
            text = el["text"].lstrip("— ").strip()
            p = right_cell.add_paragraph()

            run_dash = p.add_run("— ")
            run_dash.font.name = FONT_NAME
            run_dash.font.size = Pt(8)
            run_dash.font.color.rgb = ORANGE

            run_text = p.add_run(text)
            run_text.font.name = FONT_NAME
            run_text.font.size = Pt(8)
            run_text.font.color.rgb = BLACK

            fmt = p.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0.5)
            fmt.line_spacing = Pt(9)
            # Hanging indent for wrapped lines
            fmt.left_indent = Inches(0.2)
            fmt.first_line_indent = Inches(-0.2)

        elif t == "education_line":
            text = el["text"]
            is_sub = text.startswith("—")

            if is_sub:
                text = text.lstrip("— ").strip()
                p = right_cell.add_paragraph()
                run_dash = p.add_run("— ")
                run_dash.font.name = FONT_NAME
                run_dash.font.size = Pt(8.5)
                run_dash.font.color.rgb = ORANGE
                run_text = p.add_run(text)
                run_text.font.name = FONT_NAME
                run_text.font.size = Pt(8.5)
                run_text.font.color.rgb = BLACK
                fmt = p.paragraph_format
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(1)
            else:
                # Bold education line
                add_para(right_cell, text,
                         font_size=9, color=BLACK, bold=True,
                         space_before=1, space_after=1)

    # ── Save ──
    doc.save(output_path)
    print(f"Saved: {output_path}")


def add_separator_line(cell, color=ORANGE):
    """Add a thin colored line under a section header."""
    p = cell.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(2)
    fmt.line_spacing = Pt(2)

    pPr = p._p.get_or_add_pPr()
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}" if isinstance(color, RGBColor) else "E8530E"
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{hex_color}"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


# ── Main ─────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 2:
        print("Usage: python3 scripts/export-resume-docx.py <input.md> [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    if not os.path.exists(input_path):
        print(f"Error: {input_path} not found")
        sys.exit(1)

    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        stem = Path(input_path).stem
        output_path = str(Path.home() / "Desktop" / f"{stem}.docx")

    sidebar_lines, main_lines = parse_resume_md(input_path)
    sidebar_blocks = parse_sidebar_blocks(sidebar_lines)
    main_elements = parse_main_content(main_lines)

    print(f"Parsed: {len(sidebar_blocks)} sidebar blocks, {len(main_elements)} main elements")
    build_docx(sidebar_blocks, main_elements, output_path)


if __name__ == "__main__":
    main()

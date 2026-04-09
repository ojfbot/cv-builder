#!/usr/bin/env python3
"""
Export a cover letter markdown file to a styled .docx.

Usage:
    python3 scripts/export-cover-letter-docx.py <input.md> [output.docx]

If output path is omitted, writes to ~/Desktop/ with the input filename.
"""

import sys
import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Brand colors (matches resume export) ─────────────────────
ORANGE = RGBColor(0xE8, 0x53, 0x0E)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0x55, 0x55, 0x55)
FONT_NAME = "Calibri"

# ── Layout ───────────────────────────────────────────────────
PAGE_TOP_MARGIN = Inches(0.5)
PAGE_BOTTOM_MARGIN = Inches(0.5)
PAGE_LEFT_MARGIN = Inches(0.75)
PAGE_RIGHT_MARGIN = Inches(0.75)

BODY_FONT_SIZE = Pt(10.5)
BODY_LINE_SPACING = Pt(13.5)
PARA_SPACING = Pt(6)


def parse_cover_letter_md(filepath):
    """
    Parse markdown cover letter into paragraphs.
    Strips the H1 title line (if present) and blank lines between paragraphs.
    Returns a list of paragraph strings.
    """
    with open(filepath, "r") as f:
        text = f.read()

    lines = text.strip().split("\n")

    # Strip leading H1 title (e.g. "# Cover Letter — ...")
    if lines and lines[0].startswith("# "):
        lines = lines[1:]

    # Collapse lines into paragraphs (separated by blank lines)
    paragraphs = []
    current = []
    for line in lines:
        stripped = line.strip()
        if stripped == "":
            if current:
                paragraphs.append(" ".join(current))
                current = []
        else:
            current.append(stripped)
    if current:
        paragraphs.append(" ".join(current))

    return paragraphs


def build_docx(paragraphs, output_path):
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = PAGE_TOP_MARGIN
    section.bottom_margin = PAGE_BOTTOM_MARGIN
    section.left_margin = PAGE_LEFT_MARGIN
    section.right_margin = PAGE_RIGHT_MARGIN

    # ── Header: name + contact ──
    header_para = doc.add_paragraph()
    run = header_para.add_run("James O'Connor")
    run.font.name = FONT_NAME
    run.font.size = Pt(18)
    run.font.color.rgb = BLACK
    run.font.bold = True
    header_para.paragraph_format.space_after = Pt(2)
    header_para.paragraph_format.space_before = Pt(0)

    contact_para = doc.add_paragraph()
    contact_text = "jfo192877@gmail.com  |  832-648-8222  |  jim.software  |  github.com/ojfbot"
    run = contact_para.add_run(contact_text)
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    run.font.color.rgb = LIGHT_GRAY
    contact_para.paragraph_format.space_after = Pt(2)
    contact_para.paragraph_format.space_before = Pt(0)

    # ── Orange rule ──
    rule_para = doc.add_paragraph()
    rule_para.paragraph_format.space_before = Pt(0)
    rule_para.paragraph_format.space_after = Pt(12)
    rule_para.paragraph_format.line_spacing = Pt(2)
    pPr = rule_para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="E8530E"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # ── Body paragraphs ──
    for para_text in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(para_text)
        run.font.name = FONT_NAME
        run.font.size = BODY_FONT_SIZE
        run.font.color.rgb = DARK_GRAY

        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = PARA_SPACING
        fmt.line_spacing = BODY_LINE_SPACING

    doc.save(output_path)
    print(f"Saved: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 scripts/export-cover-letter-docx.py <input.md> [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    if not os.path.exists(input_path):
        print(f"Error: {input_path} not found")
        sys.exit(1)

    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        stem = Path(input_path).stem
        output_path = str(Path.home() / "Desktop" / f"{stem}.docx")

    paragraphs = parse_cover_letter_md(input_path)
    print(f"Parsed: {len(paragraphs)} paragraphs")
    build_docx(paragraphs, output_path)


if __name__ == "__main__":
    main()
